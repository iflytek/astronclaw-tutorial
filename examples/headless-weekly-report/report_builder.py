"""Build and verify a weekly-report DOCX inside an explicitly allowed workspace."""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import os
import tempfile
import zipfile
from collections import Counter
from datetime import date
from pathlib import Path
from typing import Any

REQUIRED_HEADERS = ("project", "status", "progress", "next_step", "owner")


def _workspace_path(workspace: Path, relative_path: str, suffix: str) -> Path:
    root = workspace.expanduser().resolve()
    if not relative_path or Path(relative_path).is_absolute():
        raise ValueError("Use a non-empty path relative to REPORT_WORKSPACE")
    candidate = (root / relative_path).resolve()
    try:
        candidate.relative_to(root)
    except ValueError as exc:
        raise ValueError("The requested path escapes REPORT_WORKSPACE") from exc
    if candidate.suffix.lower() != suffix:
        raise ValueError(f"Expected a {suffix} file: {relative_path}")
    return candidate


def _load_rows(source: Path) -> list[dict[str, str]]:
    if not source.is_file():
        raise FileNotFoundError(f"CSV source not found: {source.name}")
    with source.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.DictReader(handle)
        headers = tuple(reader.fieldnames or ())
        missing = [name for name in REQUIRED_HEADERS if name not in headers]
        if missing:
            raise ValueError("CSV is missing required columns: " + ", ".join(missing))
        rows = [
            {name: (row.get(name) or "").strip() for name in REQUIRED_HEADERS}
            for row in reader
        ]
    if not rows:
        raise ValueError("CSV must contain at least one data row")
    if any(not row["project"] for row in rows):
        raise ValueError("Every CSV row must have a project name")
    return rows


def inspect_source(workspace: Path, source_csv: str) -> dict[str, Any]:
    """Return a bounded summary without exposing files outside the workspace."""
    source = _workspace_path(workspace, source_csv, ".csv")
    rows = _load_rows(source)
    statuses = Counter(row["status"] or "Unspecified" for row in rows)
    owners = sorted({row["owner"] for row in rows if row["owner"]})
    return {
        "source": source.relative_to(workspace.expanduser().resolve()).as_posix(),
        "row_count": len(rows),
        "status_counts": dict(sorted(statuses.items())),
        "owners": owners,
        "required_columns": list(REQUIRED_HEADERS),
    }


def _verify_docx(path: Path) -> dict[str, Any]:
    if not path.is_file() or path.stat().st_size == 0:
        raise RuntimeError("DOCX output was not created or is empty")
    try:
        with zipfile.ZipFile(path) as archive:
            names = set(archive.namelist())
            required = {"[Content_Types].xml", "word/document.xml"}
            if not required.issubset(names):
                raise RuntimeError("Output is a ZIP file but not a valid DOCX package")
            bad_member = archive.testzip()
            if bad_member:
                raise RuntimeError(f"DOCX contains a corrupt member: {bad_member}")
    except zipfile.BadZipFile as exc:
        raise RuntimeError("Output is not a valid DOCX ZIP package") from exc
    digest = hashlib.sha256(path.read_bytes()).hexdigest()
    return {"size_bytes": path.stat().st_size, "sha256": digest}


def build_report(
    workspace: Path,
    source_csv: str,
    output_docx: str,
    week_ending: str,
    title: str = "Weekly project report",
    overwrite: bool = False,
) -> dict[str, Any]:
    """Create a formatted DOCX atomically and return verification metadata."""
    try:
        parsed_date = date.fromisoformat(week_ending)
    except ValueError as exc:
        raise ValueError("week_ending must use YYYY-MM-DD") from exc

    source = _workspace_path(workspace, source_csv, ".csv")
    output = _workspace_path(workspace, output_docx, ".docx")
    rows = _load_rows(source)
    if output.exists() and not overwrite:
        raise FileExistsError(
            f"Output already exists: {output.name}; set overwrite=true to replace it"
        )
    output.parent.mkdir(parents=True, exist_ok=True)

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError(
            "Install dependencies with: pip install -r requirements.txt"
        ) from exc

    document = Document()
    heading = document.add_heading(title.strip() or "Weekly project report", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(f"Week ending: {parsed_date.isoformat()}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    statuses = Counter(row["status"] or "Unspecified" for row in rows)
    document.add_heading("Executive summary", level=1)
    document.add_paragraph(
        f"{len(rows)} projects reported. "
        + "; ".join(f"{name}: {count}" for name, count in sorted(statuses.items()))
        + "."
    )

    document.add_heading("Project details", level=1)
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    labels = ("Project", "Status", "Progress", "Owner", "Next step")
    for cell, label in zip(table.rows[0].cells, labels, strict=True):
        cell.text = label
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        values = (
            row["project"],
            row["status"],
            row["progress"],
            row["owner"],
            row["next_step"],
        )
        for cell, value in zip(cells, values, strict=True):
            cell.text = value

    document.add_heading("Risks and follow-up", level=1)
    risks = [
        row for row in rows if row["status"].strip().lower() in {"blocked", "at risk"}
    ]
    if risks:
        for row in risks:
            document.add_paragraph(
                f"{row['project']}: {row['status']} — {row['next_step']}",
                style="List Bullet",
            )
    else:
        document.add_paragraph("No blocked or at-risk projects were reported.")

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10.5)

    temporary_path: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(
            prefix=f".{output.stem}-", suffix=".docx", dir=output.parent, delete=False
        ) as handle:
            temporary_path = Path(handle.name)
        document.save(temporary_path)
        verification = _verify_docx(temporary_path)
        if output.exists() and not overwrite:
            raise FileExistsError(
                f"Output appeared during generation: {output.name}; retry with a new name"
            )
        os.replace(temporary_path, output)
        temporary_path = None
    finally:
        if temporary_path and temporary_path.exists():
            temporary_path.unlink()

    return {
        "output": output.relative_to(workspace.expanduser().resolve()).as_posix(),
        "week_ending": parsed_date.isoformat(),
        "row_count": len(rows),
        **verification,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--workspace", type=Path, required=True)
    parser.add_argument("--source", default="data.csv")
    parser.add_argument("--output", required=True)
    parser.add_argument("--week-ending", required=True)
    parser.add_argument("--title", default="Weekly project report")
    parser.add_argument("--overwrite", action="store_true")
    args = parser.parse_args()
    result = build_report(
        args.workspace,
        args.source,
        args.output,
        args.week_ending,
        args.title,
        args.overwrite,
    )
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
